from copy import deepcopy

from docx import Document


SOURCE = "EverWind_Server_Requirements_Draft.docx"
OUTPUT = "EverWind_Server_Requirements_Draft_QuestAdded.docx"


QUEST_ROWS = {
    0: [
        [
            "FR-16",
            "퀘스트 수주 및 상태 조회",
            "서버는 사용자가 수락 가능한 퀘스트 목록을 조회하고 퀘스트를 수주할 수 있도록 지원해야 한다.",
            "사용자 식별값, NPC ID 또는 퀘스트 ID",
            "사용자의 선행 조건과 현재 진행 상태를 확인한 뒤 수락 가능한 퀘스트 목록을 반환하거나 선택된 퀘스트를 활성 상태로 저장한다.",
            "수락 가능 퀘스트 목록, 수주 결과, 활성 퀘스트 상태",
            "상",
        ],
        [
            "FR-17",
            "퀘스트 진행도 반영",
            "서버는 몬스터 처치, 아이템 수집, 대화 완료 등 퀘스트 조건 달성 여부를 추적해야 한다.",
            "사용자 행동 이벤트, 퀘스트 목표 정보",
            "수신한 게임 이벤트를 활성 퀘스트 조건과 비교하여 진행도를 갱신하고 변경 사항을 저장한다.",
            "갱신된 퀘스트 진행도, 목표 달성 여부",
            "상",
        ],
        [
            "FR-18",
            "퀘스트 완료 및 보상 지급",
            "서버는 완료 조건을 만족한 퀘스트를 종료 처리하고 보상을 지급해야 한다.",
            "사용자 식별값, 완료 요청 퀘스트 ID",
            "퀘스트 완료 조건을 재검증한 뒤 경험치, 아이템, 재화 등의 보상을 지급하고 완료 상태를 저장한다.",
            "퀘스트 완료 결과, 지급 보상 내역, 갱신된 사용자 상태",
            "상",
        ],
    ],
    1: [
        [
            "PR-09",
            "퀘스트 수주 처리 절차",
            "사용자 식별값, NPC ID 또는 퀘스트 ID",
            "1. 사용자의 레벨, 선행 퀘스트, 반복 가능 여부를 확인한다. / 2. 수주 가능한 퀘스트 목록을 조회한다. / 3. 사용자가 선택한 퀘스트를 활성 상태로 등록한다. / 4. 초기 진행도 정보를 저장하고 클라이언트에 전달한다.",
            "활성 퀘스트 목록과 수주 결과가 전달된다.",
        ],
        [
            "PR-10",
            "퀘스트 진행도 갱신 절차",
            "몬스터 처치, 아이템 획득, 상호작용 완료 이벤트",
            "1. 이벤트와 연관된 활성 퀘스트를 조회한다. / 2. 퀘스트 목표 조건과 이벤트 내용을 비교한다. / 3. 충족된 목표의 진행도를 증가시킨다. / 4. 목표 완료 여부와 전체 퀘스트 완료 가능 여부를 판정한다. / 5. 변경된 진행도를 저장하고 클라이언트에 전송한다.",
            "갱신된 진행도와 완료 가능 상태가 전달된다.",
        ],
        [
            "PR-11",
            "퀘스트 완료 및 보상 지급 절차",
            "사용자 식별값, 완료 요청 퀘스트 ID",
            "1. 대상 퀘스트가 활성 상태인지 확인한다. / 2. 모든 완료 조건 충족 여부를 재검증한다. / 3. 보상 테이블을 조회한다. / 4. 경험치, 아이템, 재화를 사용자 데이터에 반영한다. / 5. 퀘스트를 완료 상태로 변경하고 이력을 저장한다. / 6. 완료 결과와 보상 정보를 클라이언트에 응답한다.",
            "퀘스트 완료 결과와 보상 지급 내역이 전달된다.",
        ],
    ],
    2: [
        [
            "IF-16",
            "퀘스트 목록 조회/수주",
            "클라이언트 ↔ 서버",
            "C2S_QUEST_LIST_REQ / S2C_QUEST_LIST_ACK / C2S_QUEST_ACCEPT_REQ / S2C_QUEST_ACCEPT_ACK",
            "사용자 ID, NPC ID 또는 퀘스트 ID / 퀘스트 목록, 수주 결과, 진행 상태",
            "수락 가능한 퀘스트 조회 및 퀘스트 수주 처리",
        ],
        [
            "IF-17",
            "퀘스트 진행도 동기화",
            "서버 → 클라이언트",
            "S2C_QUEST_PROGRESS_ACK",
            "퀘스트 ID, 목표 ID, 현재 진행도, 목표 달성 여부",
            "게임 이벤트 반영 후 퀘스트 진행 상태를 갱신하여 전송",
        ],
        [
            "IF-18",
            "퀘스트 완료/보상 수령",
            "클라이언트 ↔ 서버",
            "C2S_QUEST_COMPLETE_REQ / S2C_QUEST_COMPLETE_ACK",
            "퀘스트 ID / 완료 결과, 경험치, 아이템, 재화",
            "퀘스트 완료 처리 및 보상 지급 결과 반환",
        ],
        [
            "IF-19",
            "DB 퀘스트 상태 조회/저장",
            "서버 ↔ DB",
            "QuestMaster, UserQuest, QuestReward",
            "퀘스트 기본 정보, 사용자 진행 상태, 보상 정보",
            "퀘스트 마스터 데이터 조회와 사용자 퀘스트 진행 상태 저장",
        ],
    ],
}


def clear_paragraph(paragraph):
    element = paragraph._element
    for child in list(element):
        element.remove(child)


def replace_cell_text(cell, text):
    paragraph = cell.paragraphs[0]
    runs = paragraph.runs
    if runs:
        runs[0].text = text
        for run in runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)

    for extra_paragraph in cell.paragraphs[1:]:
        clear_paragraph(extra_paragraph)


def append_styled_row(table, values):
    template_row = table.rows[-1]
    new_tr = deepcopy(template_row._tr)
    table._tbl.append(new_tr)
    new_row = table.rows[-1]
    for cell, value in zip(new_row.cells, values):
        replace_cell_text(cell, value)


def main():
    doc = Document(SOURCE)
    for table_index, rows in QUEST_ROWS.items():
        for row in rows:
            append_styled_row(doc.tables[table_index], row)
    doc.save(OUTPUT)


if __name__ == "__main__":
    main()
